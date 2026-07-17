"""
Word document parser — extracts paragraphs and embedded tables.
Owner: Member 2 — ML & Document Intelligence Lead
"""
from docx import Document
from loguru import logger


class WordParser:
    """Parse .docx files into plain text preserving tables."""

    def extract(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            parts = []

            # Paragraphs (preserves heading hierarchy via style names)
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name if para.style else ""
                if "Heading" in style:
                    parts.append(f"\n## {text}")
                else:
                    parts.append(text)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            return "\n".join(parts)

        except Exception as e:
            logger.error(f"DOCX parse failed for {file_path}: {e}")
            return ""
